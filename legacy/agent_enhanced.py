"""
Workflow Thief Arena - 智能 Agent
包含完整文档生成功能，真实网络搜索 + LLM 分析！
"""
import json
import sys
import time
import os
from datetime import datetime
from typing import Dict, List, Optional, Callable
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from workflow_data import get_preset

# 进度回调函数
progress_callback: Optional[Callable] = None


def set_progress_callback(callback: Optional[Callable]):
    """设置进度回调函数"""
    global progress_callback
    progress_callback = callback


def update_progress(stage: str, progress: int):
    """更新进度"""
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {stage}")
    if progress_callback:
        progress_callback(stage, progress)


def get_evidence_urls(keyword: str) -> List[Dict]:
    """使用真实 Bing 搜索获取证据链！"""
    update_progress("正在搜索网络...", 10)
    from bing_search import search_bing
    results = search_bing(keyword, num_results=8)

    # 补充网页内容
    try:
        update_progress("正在抓取网页内容...", 20)
        from web_scraper import enrich_search_results
        results = enrich_search_results(results)
    except Exception as e:
        print(f"[Agent] 网页抓取失败: {e}")

    return results


# 保留别名，兼容旧代码
get_preset_evidence = get_evidence_urls


def analyze_with_llm(search_results: List[Dict], keyword: str) -> Dict:
    """
    使用 LLM 进行真实分析

    Args:
        search_results: 搜索结果
        keyword: 关键词

    Returns:
        分析结果字典
    """
    try:
        update_progress("正在调用 LLM 分析...", 30)
        from llm_client import analyze_workflow_with_llm
        print(f"[Agent] 正在调用火山引擎 API 分析: {keyword}")
        result = analyze_workflow_with_llm(search_results, keyword)
        print(f"[Agent] LLM 分析完成！")
        return result
    except Exception as e:
        print(f"[Agent] LLM 分析失败: {e}，使用预设数据")
        return get_preset(keyword)


def sanitize_mermaid_text(text: str) -> str:
    """清理 Mermaid 文本"""
    if not text:
        return ""
    text = str(text)
    text = text.replace('"', "'")
    text = text.replace('\n', ' ')
    text = text.replace('\\n', ' ')
    text = text.replace('，', ',')
    text = text.replace('（', ' ')
    text = text.replace('）', ' ')
    text = text.replace('(', ' ')
    text = text.replace(')', ' ')
    if len(text) > 20:
        text = text[:17] + "..."
    return text.strip()


def generate_mermaid_original_flow(data: Dict) -> str:
    """生成原始工作流 Mermaid"""
    trigger = sanitize_mermaid_text(data.get('trigger_condition', data.get('trigger', '触发条件')))
    if len(trigger) > 15:
        trigger = trigger[:12] + "..."

    mermaid = "flowchart TD\n"
    mermaid += "    Start((开始)) --> Roles[参与角色]\n"
    mermaid += f"    Roles --> Trigger[{trigger}]\n"

    prev_step = "Trigger"
    for step in data.get('steps', []):
        step_id = f"Step{step['step_number']}"
        desc = sanitize_mermaid_text(step['description'])
        label = f"{step['step_number']}.{desc}"

        if step.get('is_decision', False):
            mermaid += f"    {prev_step} --> {step_id}{{{label}}}\n"
            mermaid += f"    {step_id} --> Next{step['step_number']}[下一步]\n"
            mermaid += f"    {step_id} --> Reject{step['step_number']}[退回]\n"
            prev_step = f"Next{step['step_number']}"
        else:
            mermaid += f"    {prev_step} --> {step_id}[{label}]\n"
            prev_step = step_id

    mermaid += f"    {prev_step} --> End((结束))\n"
    return mermaid


def generate_mermaid_agent_flow(data: Dict) -> str:
    """生成 Agent 改造后的 Mermaid"""
    agent_flow = data.get('agent_flow', {})
    intervention_map = {p['step_number']: p for p in agent_flow.get('intervention_points', [])}

    trigger = sanitize_mermaid_text(data.get('trigger_condition', data.get('trigger', '触发条件')))
    if len(trigger) > 15:
        trigger = trigger[:12] + "..."

    mermaid = "flowchart TD\n"
    mermaid += "    Start((开始)) --> Roles[参与角色]\n"
    mermaid += f"    Roles --> Trigger[{trigger}]\n"

    prev_step = "Trigger"
    for step in data.get('steps', []):
        step_id = f"Step{step['step_number']}"
        intervention = intervention_map.get(step['step_number'])
        desc = sanitize_mermaid_text(step['description'])
        base_label = f"{step['step_number']}.{desc}"

        if intervention:
            int_type = intervention['intervention_type']
            if "完全自动" in int_type:
                label = f"{base_label}-自动"
            elif "辅助" in int_type:
                label = f"{base_label}-Agent辅助"
            else:
                label = f"{base_label}-人工确认"
            mermaid += f"    {prev_step} --> {step_id}[{label}]\n"

            if intervention.get('risk_control'):
                risk_id = f"Risk{step['step_number']}"
                mermaid += f"    {step_id} --> {risk_id}[风控]\n"
                prev_step = risk_id
            else:
                prev_step = step_id
        else:
            label = f"{base_label}-人工"
            mermaid += f"    {prev_step} --> {step_id}[{label}]\n"
            prev_step = step_id

    mermaid += f"    {prev_step} --> End((结束))\n"
    return mermaid


def generate_word_document(data: Dict, output_path: str):
    """生成 Word 文档"""
    doc = Document()

    title = doc.add_heading('Workflow Thief Arena - 工作流分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    workflow_name = data.get('workflow_name', data.get('name', '未知'))
    doc.add_paragraph(f'分析主题: {workflow_name}')
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_page_break()

    doc.add_heading('1. 项目概述', level=1)
    p = doc.add_paragraph()
    p.add_run('项目名称: ').bold = True
    p.add_run(f'Workflow Thief Arena - {workflow_name} 智能分析\n')
    p.add_run('所属行业: ').bold = True
    p.add_run(f'{data.get("industry", "未知")}\n')
    p.add_run('项目目标: ').bold = True
    p.add_run('通过智能 Agent 技术，分析现有工作流程的低效点，提供自动化改造方案，提升业务效率。\n')
    p.add_run('核心价值: ').bold = True

    cost = data.get('cost', {})
    savings = cost.get('savings', {})
    roi = cost.get('roi', {})
    annual_savings = savings.get('annual_savings', 0)
    first_year_roi = roi.get('first_year_roi', 0)

    p.add_run(f'预计年度节省: {annual_savings} 元，ROI: {first_year_roi}%\n')

    doc.add_heading('2. Agent 搭建说明', level=1)
    doc.add_paragraph('本项目使用以下技术栈:', style='List Number')
    doc.add_paragraph('搜索: Bing 真实网络搜索')
    doc.add_paragraph('网页抓取: requests + BeautifulSoup')
    doc.add_paragraph('LLM: Anthropic Claude (火山引擎)')
    doc.add_paragraph('文档生成: python-docx')

    doc.add_heading('3. 使用的技术与工具', level=1)
    tech_table = doc.add_table(rows=1, cols=2)
    tech_table.style = 'Table Grid'
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text = '类别'
    hdr_cells[1].text = '技术/工具'
    tech_data = [
        ('开发语言', 'Python 3'),
        ('网络搜索', 'Bing Search'),
        ('网页抓取', 'requests + BeautifulSoup'),
        ('后端框架', 'Flask'),
        ('文档生成', 'python-docx'),
        ('流程图', 'Mermaid'),
    ]
    for cat, tech in tech_data:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = tech

    doc.add_heading('4. 设计理念', level=1)
    doc.add_paragraph('真实数据驱动: 基于网络搜索的真实证据链', style='List Number')
    doc.add_paragraph('ROI 导向: 明确成本收益分析')
    doc.add_paragraph('7 天 MVP 验证: 快速验证可行性')
    doc.add_paragraph('人机协作: 清晰区分 Agent 和人类责任')

    doc.add_heading('5. 公开证据链', level=1)
    doc.add_paragraph('以下是本分析报告的信息来源:')
    for i, evidence in enumerate(data.get('evidence_urls', []), 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{evidence.get("title", "无标题")}\n')
        p.add_run(f'链接: {evidence.get("url", "无链接")}')

    doc.add_heading('6. 原始工作流流程图', level=1)
    mermaid_original = generate_mermaid_original_flow(data)
    code_para = doc.add_paragraph()
    code_para.add_run(mermaid_original).font.name = 'Courier New'

    doc.add_heading('7. Agent 改造后的新流程图', level=1)
    mermaid_agent = generate_mermaid_agent_flow(data)
    code_para = doc.add_paragraph()
    code_para.add_run(mermaid_agent).font.name = 'Courier New'

    doc.add_paragraph('图例说明:')
    legend_table = doc.add_table(rows=1, cols=2)
    legend_table.style = 'Table Grid'
    hdr_cells = legend_table.rows[0].cells
    hdr_cells[0].text = '类型'
    hdr_cells[1].text = '说明'
    legend_data = [
        ('-自动', 'Agent 完全自动执行'),
        ('-Agent 辅助', 'Agent 辅助人类执行'),
        ('-人工确认', '需要人类确认'),
        ('-人工', '保持由人类执行'),
    ]
    for typ, desc in legend_data:
        row_cells = legend_table.add_row().cells
        row_cells[0].text = typ
        row_cells[1].text = desc

    doc.add_page_break()
    doc.add_heading('8. 低效点与 Agent 介入点', level=1)
    doc.add_heading('8.1 低效点分析', level=2)
    pain_points = data.get('pain_points', {})
    pain_labels = {
        'manual_repetition': '人工重复',
        'information_movement': '信息搬运',
        'judgment_cost': '判断成本',
        'communication_cost': '沟通成本',
        'waiting_cost': '等待成本',
        'audit_cost': '审核成本'
    }
    for key, pain in pain_points.items():
        if isinstance(pain, dict):
            doc.add_paragraph(f"{pain_labels.get(key, key)}:")
            doc.add_paragraph(f"描述: {pain.get('description', '')}")
            doc.add_paragraph(f"占比: {pain.get('time_percentage', 0)}%")

    doc.add_heading('8.2 Agent 介入点', level=2)
    agent_flow = data.get('agent_flow', {})
    for point in agent_flow.get('intervention_points', []):
        doc.add_paragraph(f"步骤{point.get('step_number', '')}: {point.get('intervention_type', '')}")
        doc.add_paragraph(f"描述: {point.get('description', '')}")

    doc.add_page_break()
    doc.add_heading('9. 产品方案', level=1)
    product = agent_flow.get('product_solution', {})
    doc.add_paragraph(f'产品名称: {product.get("product_name", "智能流程助手")}')

    doc.add_heading('核心功能', level=2)
    for feat in product.get('core_features', []):
        doc.add_paragraph(feat, style='List Bullet')

    doc.add_page_break()
    doc.add_heading('10. 7 天 MVP 验证计划', level=1)
    mvp_plan = agent_flow.get('mvp_plan', [])
    for plan in mvp_plan:
        if isinstance(plan, dict):
            doc.add_paragraph(f"Day {plan.get('day', '')}: {plan.get('phase', '')}")
            doc.add_paragraph(f"任务: {plan.get('task', '')}")

    doc.add_page_break()
    doc.add_heading('11. 成本与效率记录', level=1)
    doc.add_paragraph(f'月度节省: {savings.get("monthly_savings", 0)} 元')
    doc.add_paragraph(f'年度节省: {annual_savings} 元')
    doc.add_paragraph(f'回本周期: {roi.get("break_even_months", 0)} 个月')
    doc.add_paragraph(f'第一年 ROI: {first_year_roi}%')

    doc.save(output_path)
    print(f"[文档] Word 文档已生成: {output_path}")


def run_full_agent(keyword: str, output_dir: str = "output", use_real_llm: bool = True) -> Optional[Dict]:
    """
    运行完整 Agent 流程

    Args:
        keyword: 搜索关键词
        output_dir: 输出目录
        use_real_llm: 是否使用真实 LLM 分析（False 则使用预设数据）

    Returns:
        分析结果字典
    """
    os.makedirs(output_dir, exist_ok=True)

    print(f"\n{'='*80}")
    print("  WORKFLOW THIEF ARENA - 智能 Agent 启动")
    print(f"{'='*80}")
    print(f"开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"分析目标: {keyword}")
    print(f"模式: {'真实 LLM 分析' if use_real_llm else '预设数据演示'}")
    print(f"{'='*80}")

    update_progress("正在初始化...", 5)

    # 真实网络搜索获取证据链
    print(f"\n[{datetime.now().strftime('%H:%M:%S')}] 网络搜索证据链...")
    evidence_urls = get_evidence_urls(keyword)

    if use_real_llm:
        # 使用 LLM 进行真实分析
        print(f"[{datetime.now().strftime('%H:%M:%S')}] LLM 深度分析...")
        result = analyze_with_llm(evidence_urls, keyword)
    else:
        # 使用预设数据
        update_progress("正在加载预设数据...", 40)
        print(f"[{datetime.now().strftime('%H:%M:%S')}] 加载预设数据...")
        result = get_preset(keyword) or get_preset('recruitment')
        result['evidence_urls'] = evidence_urls

    if not result:
        update_progress("分析失败", 0)
        print(f"[{datetime.now().strftime('%H:%M:%S')}] [错误] 无法获取数据")
        return None

    update_progress("正在生成文档...", 70)

    # 补充元数据
    result['generated_at'] = datetime.now().isoformat()
    result['keyword'] = keyword

    # 生成文档
    print(f"\n[{datetime.now().strftime('%H:%M:%S')}] 生成分析文档...")
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_keyword = keyword.replace(' ', '_').replace('/', '_')

    # 保存 JSON
    json_path = os.path.join(output_dir, f"{safe_keyword}_{timestamp}.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"[{datetime.now().strftime('%H:%M:%S')}] [保存] JSON: {json_path}")

    update_progress("正在生成 Word 文档...", 80)

    # 保存 Word
    docx_path = os.path.join(output_dir, f"{safe_keyword}_{timestamp}.docx")
    generate_word_document(result, docx_path)

    update_progress("正在生成 Mermaid 流程图...", 90)

    # 保存 Mermaid
    mermaid_path = os.path.join(output_dir, f"{safe_keyword}_{timestamp}_mermaid.txt")
    with open(mermaid_path, 'w', encoding='utf-8') as f:
        f.write("=== 原始工作流 ===\n")
        f.write(generate_mermaid_original_flow(result))
        f.write("\n\n=== Agent 改造后 ===\n")
        f.write(generate_mermaid_agent_flow(result))
    print(f"[{datetime.now().strftime('%H:%M:%S')}] [保存] Mermaid: {mermaid_path}")

    update_progress("分析完成！", 100)

    print(f"\n{'='*80}")
    print("  分析完成！")
    print(f"{'='*80}")
    print(f"输出目录: {os.path.abspath(output_dir)}")
    print(f"JSON: {os.path.basename(json_path)}")
    print(f"Word: {os.path.basename(docx_path)}")
    print(f"{'='*80}")

    return result


if __name__ == '__main__':
    use_real_llm = True

    if len(sys.argv) > 1:
        if sys.argv[1] == '--preset':
            use_real_llm = False
            keyword = sys.argv[2] if len(sys.argv) > 2 else '招聘筛选流程'
        else:
            keyword = sys.argv[1]
    else:
        keyword = '招聘筛选流程'

    run_full_agent(keyword, use_real_llm=use_real_llm)
